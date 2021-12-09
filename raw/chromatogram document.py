from docx import Document
from docx.shared import Inches

document = Document()

paragraph = document.add_paragraph()
paragraph.add_run('Table 1. Benthiavalicarb-isopropyl+Oxathiapiprolin WG 중 Benthiavalicarb-isopropyl 분석근거 및 결과 (시작)').bold = True

table = document.add_table(rows=7, cols=6)
print (type(table.style))

table.style = 'Table Grid'

hdr1_cells = table.rows[0].cells
hdr1_cells[1].text = '무게(g)'
hdr1_cells[2].text = '순도(%)'
hdr1_cells[3].text = 'A.I area'
hdr1_cells[4].text = 'I.S area'
hdr1_cells[5].text = '결과'

hdr2_cells = table.rows[1].cells
hdr2_cells[0].text = 'Factor'

document.add_paragraph('계산식').bold=True

document.add_page_break()

document.save('demo.docx')